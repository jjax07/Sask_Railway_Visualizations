#!/usr/bin/env python3
"""Create a Word document with Saskatchewan Railway data issues for consultation."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Saskatchewan Railway Visualization - Data Issues', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Introduction
doc.add_paragraph(
    'The following data issues have been identified during the development of the '
    'Saskatchewan Railway Visualization project. Input from historical researchers '
    'and railway experts would be valuable in resolving these issues.'
)

# Issue 1: Missing Railway Arrival Years
doc.add_heading('1. Missing Railway Arrival Years (28 entries)', level=1)

doc.add_paragraph(
    'When processing settlement-railway pairs, approximately 28 entries could not be '
    'matched to arrival years. These need historical research to determine when each '
    'railway reached these settlements.'
)

doc.add_heading('What is needed:', level=2)
doc.add_paragraph(
    'For each settlement-railway pair listed below, we need the year when the railway '
    'first provided service to that settlement. Sources might include:'
)

bullets = doc.add_paragraph()
bullets.add_run('• Railway company historical records\n')
bullets.add_run('• Saskatchewan Archives\n')
bullets.add_run('• Local newspaper archives (arrival announcements)\n')
bullets.add_run('• Municipal histories\n')
bullets.add_run('• Canadian Railway Historical Association records')

doc.add_paragraph(
    'Note: The specific list of 28 unmatched settlement-railway pairs can be generated '
    'by running the update_multi_railways.py script and reviewing the output.'
)

# Issue 2: CPR Takeover of QLSRSC Line
doc.add_heading('2. CPR Takeover of QLSRSC Line', level=1)

doc.add_paragraph(
    "The QLSRSC (Qu'Appelle, Long Lake & Saskatchewan Railway) was leased by CPR "
    "from approximately 1896-1906. Settlements on this line should potentially show "
    "both QLSRSC and CPR service during this period."
)

doc.add_heading('Affected Settlements:', level=2)
doc.add_paragraph(
    'Settlements between Regina and Prince Albert via Saskatoon, including:'
)

affected = doc.add_paragraph()
affected.add_run(
    'Lumsden, Craven, Bethune, Chamberlain, Craik, Davidson, Hanley, Dundurn, '
    'Saskatoon, Warman, Osler, Rosthern, Duck Lake, and others along this route.'
)

doc.add_heading('Questions for researchers:', level=2)
questions = doc.add_paragraph()
questions.add_run('• What were the exact dates of the CPR lease of QLSRSC?\n')
questions.add_run('• Did CPR operate the line under the QLSRSC name, or rebrand it?\n')
questions.add_run('• Were there operational differences during the lease period?\n')
questions.add_run('• When did CPR formally acquire (vs. lease) the line?\n')
questions.add_run('• Should settlements show both "QLSRSC" and "CPR" during the lease, or just one?')

# Issue 3: Post-1918 CN Consolidation
doc.add_heading('3. Post-1918 CN Consolidation', level=1)

doc.add_paragraph(
    'The Canadian Northern Railway (CNoR) and Grand Trunk Pacific Railway (GTPR) were '
    'absorbed into Canadian National (CN) between 1918-1923. The current visualization '
    'ends at 1920, so this transition is only partially captured.'
)

doc.add_heading('Current handling:', level=2)
doc.add_paragraph(
    '"CN" entries in the source spreadsheet are currently mapped to the "Other" category, '
    'which may not accurately reflect the historical significance of CN\'s formation.'
)

doc.add_heading('Questions for researchers:', level=2)
questions2 = doc.add_paragraph()
questions2.add_run('• What are the key dates for CNoR absorption into CN?\n')
questions2.add_run('• What are the key dates for GTPR absorption into CN?\n')
questions2.add_run('• Should the visualization timeline extend to 1923 to show full consolidation?\n')
questions2.add_run('• How should we visually represent this transition (gradual or sudden)?')

# Contact section
doc.add_heading('How to Provide Input', level=1)
doc.add_paragraph(
    'Please provide any information, sources, or corrections to help resolve these issues. '
    'Even partial information (e.g., a single settlement\'s railway arrival year) is valuable.'
)

# Footer
doc.add_paragraph()
doc.add_paragraph('Document generated: January 2026').italic = True

# Save
output_path = '/Users/baebot/Documents/ClaudeCode/Sask_Railway_Visualizations/Saskatchewan_Railway_Data_Issues.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
